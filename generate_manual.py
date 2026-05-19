"""Generate Scholara_User_Manual.docx — run with: python generate_manual.py"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────
RED  = RGBColor(0xC8, 0x43, 0x0B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(22)
    p.paragraph_format.space_after    = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size      = Pt(17)
    run.font.color.rgb = RED
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(14)
    p.paragraph_format.space_after    = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size      = Pt(12)
    run.font.color.rgb = DARK
    return p


def body(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size      = Pt(10)
    run.font.color.rgb = DARK
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.2)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = DARK
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name      = 'Courier New'
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0x14, 0x14, 0x14)
    return p


def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run("TIP:  " + text)
    run.font.size      = Pt(10)
    run.font.italic    = True
    run.font.color.rgb = GREY
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("NOTE:  " + text)
    run.font.size      = Pt(10)
    run.font.italic    = True
    run.font.color.rgb = GREY
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        hdr[i].paragraphs[0].paragraph_format.space_before = Pt(2)
        hdr[i].paragraphs[0].paragraph_format.space_after  = Pt(2)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row[ci].text = cell_text
            for run in row[ci].paragraphs[0].runs:
                run.font.size = Pt(9)
            row[ci].paragraphs[0].paragraph_format.space_before = Pt(2)
            row[ci].paragraphs[0].paragraph_format.space_after  = Pt(2)
    if col_widths:
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Scholara")
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Open Knowledge, Everywhere")
r.italic = True; r.font.size = Pt(14); r.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("User Manual & Functional Verification Guide")
r.font.size = Pt(13); r.font.color.rgb = DARK

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0  ·  May 2026")
r.font.size = Pt(11); r.font.color.rgb = GREY

page_break()

# ── About ─────────────────────────────────────────────────────────
heading1("About This Guide")
body("This manual walks you through every feature of Scholara step by step. Each section includes an easy example you can follow right now and a clear 'What to verify' checklist.")
body("You do not need any technical background. All you need is Scholara running on your computer and a web browser.")

heading2("What is Scholara?")
body("Scholara is a free, open-source academic research platform designed for researchers and students in Africa and the Global South. It lets you:")
bullet("Search millions of open-access papers, books, and journals with a single query.")
bullet("Download academic PDFs and documents directly to your computer, for free.")
bullet("Convert files between formats (PDF → Word, Word → PDF, video → audio).")
bullet("Organise your research with history tracking, tags, and named collections.")
bullet("Export citations in BibTeX, RIS, and APA formats.")
bullet("Use the interactive AI demo to explain, summarise, and chat with any document.")
bullet("Use the interface in English or French.")
bullet("Install as an app on your phone or desktop (PWA).")

add_table(
    ["Component", "Technology"],
    [
        ["Backend", "Python 3.14 + FastAPI"],
        ["Frontend", "Vanilla HTML/CSS/JS — ES modules, no framework"],
        ["Database", "SQLite (scholara.db)"],
        ["AI Router", "Ollama (local) → DeepSeek (cloud) → Keyword fallback"],
        ["Download", "Direct HTTP — open-access domains only"],
        ["Conversion", "Pure-Python: pdf2docx, pypdf, mammoth, ffmpeg"],
        ["Auth (multi-user)", "python-jose[cryptography] · passlib[bcrypt]"],
    ],
    col_widths=[2.0, 4.5]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1 — GETTING STARTED
# ══════════════════════════════════════════════════════════════════
heading1("Section 1 — Getting Started")

heading2("1.1  Starting the Application")
body("Open a terminal, navigate to the Scholara folder, and run:")
code("# First time only\npython3 -m venv .venv\nsource .venv/bin/activate\npip install -r requirements.txt\n\n# Every time\npython app.py")
body("Scholara opens automatically at: http://127.0.0.1:7860")

heading2("1.2  The Status Banner")
code("=== Scholara Status ===\n  [OK] Ollama (mistral)\n  [--] DeepSeek API\n  [OK] ffmpeg (conversion)\n  [OK] Direct HTTP: always available\n======================")
tip("If Ollama shows [--], the AI router falls back to keyword matching. Search still works.")

heading2("1.3  The Interface at a Glance")
add_table(
    ["Area", "What it does"],
    [
        ["Header bar", "Logo, language toggle (EN / FR), status badges, Sign Out button"],
        ["Left sidebar tabs", "Search · URL · Convert · History · Collections · Admin"],
        ["Results grid", "Cards showing search results with Download buttons"],
        ["Right viewer", "Inline preview of PDFs, text, audio, and video files"],
        ["Demo sidebar", "Slides in from the right — AI explanation, summary, chat, slides, flowchart"],
        ["Settings modal", "API keys, download folder, dark mode, language preference"],
    ],
    col_widths=[2.2, 4.3]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2 — NATURAL LANGUAGE SEARCH
# ══════════════════════════════════════════════════════════════════
heading1("Section 2 — Natural Language Search")
body("Type a question or topic in plain English or French. The AI router selects the most relevant sources automatically. Results are deduplicated and reranked by relevance.")

heading2("2.1  English Search — Example")
bullet("Click the Search tab in the left sidebar.")
bullet("Type: machine learning applications in African agriculture")
bullet("Press Enter or click Search.")
body("What to verify:", bold=True)
bullet("Results appear within 10 seconds.")
bullet("At least one result from arXiv (📄) or OpenAlex (🔬).")
bullet("Each card shows: Title, Authors, Source badge, Download button.")

heading2("2.2  French Search — Example")
code("histoire de l’éducation en Afrique subsaharienne")
body("What to verify:", bold=True)
bullet("HAL (🏛️) results appear — France's national open archive.")
bullet("Routing label shows 'hal' and/or 'persee' as selected sources.")
tip("Scholara detects French queries and automatically prioritises HAL, Persée, OpenEdition, and Érudit.")

heading2("2.3  Search Intelligence")
add_table(
    ["Feature", "How it works"],
    [
        ["Deduplication", "Removes near-duplicate results — exact DOI match or >90% title similarity"],
        ["Reranking", "+3 language match · +2 DOI present · +2 abstract · +1 recent · +1 open-access"],
    ],
    col_widths=[2.0, 4.5]
)

heading2("2.4  Core and Global North Sources")
add_table(
    ["Source", "Language", "Always active?"],
    [
        ["arXiv", "EN", "Yes"],
        ["DOAJ", "EN/FR", "Yes"],
        ["OpenAlex", "EN", "Yes"],
        ["Internet Archive", "EN", "Yes"],
        ["Project Gutenberg", "EN", "Yes"],
        ["HAL", "FR/EN", "Yes"],
        ["Persée", "FR", "Yes"],
        ["OpenEdition", "FR/EN", "Yes"],
        ["Érudit", "FR/EN", "Yes"],
        ["Semantic Scholar", "EN", "Global North only"],
        ["PubMed", "EN", "Global North only"],
        ["CrossRef", "EN", "Global North only"],
        ["CORE", "EN", "Global North only"],
        ["BASE", "Multilingual", "Global North only"],
    ],
    col_widths=[2.0, 1.2, 2.0]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3 — DOWNLOADING FILES
# ══════════════════════════════════════════════════════════════════
heading1("Section 3 — Downloading Files")

heading2("3.1  Download from Search Results")
bullet("Run any search — click 'Download PDF' on any result card.")
body("What to verify:", bold=True)
bullet("Progress bar appears → file appears in History and file list when done.")
bullet("Clicking the file opens it in the viewer.")

heading2("3.2  Download by Pasting a URL")
bullet("Click the URL tab → paste an open-access URL:")
code("https://arxiv.org/pdf/1706.03762.pdf")
bullet("Click Download.")
body("What to verify:", bold=True)
bullet("Status: queued → running → done.")
bullet("File appears in file list and History.")

heading2("3.3  Blocked URL Test")
bullet("Paste: https://www.google.com/file.pdf → Click Download.")
body("What to verify:", bold=True)
bullet("Error message appears — no file created. Only whitelisted open-access domains are allowed.")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4 — INTERACTIVE AI DEMO
# ══════════════════════════════════════════════════════════════════
heading1("Section 4 — Interactive AI Demo")
body("The Interactive Demo lets you analyse any downloaded file or selected text with five AI-powered actions. Right-click a file entry, or select text anywhere on the page to trigger it.")

heading2("4.1  How to Trigger the Demo")
add_table(
    ["Trigger", "How"],
    [
        ["Right-click a file", "Right-click any entry in the file list or History tab"],
        ["Text selection", "Select more than 20 characters anywhere → context menu appears"],
    ],
    col_widths=[2.2, 4.3]
)

heading2("4.2  Five AI Actions")
add_table(
    ["Action", "What it produces"],
    [
        ["💡 Text Explanation", "Plain-language explanation for a non-specialist reader"],
        ["📄 Summary", "Structured: main topic · key findings · methodology · limitations"],
        ["🤖 Interactive Demo", "Multi-turn AI chat grounded in the document text"],
        ["📊 Presentation", "Slide-by-slide outline rendered as cards in the sidebar"],
        ["🔀 Flowchart", "Mermaid.js diagram of the main process or argument"],
    ],
    col_widths=[2.2, 4.3]
)

heading2("4.3  Text Explanation — Example")
bullet("Download arXiv paper 1706.03762.pdf (Section 3.2).")
bullet("Right-click the file entry → select '💡 Text Explanation'.")
body("What to verify:", bold=True)
bullet("Demo sidebar slides in from the right.")
bullet("Loading spinner appears briefly, then the explanation renders.")
bullet("Response language matches your current interface language (EN or FR).")

heading2("4.4  Interactive Chat — Example")
bullet("Right-click any PDF → '🤖 Interactive Demo'.")
bullet("Sidebar opens with an introductory response and a chat input.")
bullet("Type a follow-up question and press Enter:")
code("What is the main contribution of this paper?")
body("What to verify:", bold=True)
bullet("Response appears, grounded in the document.")
bullet("Up to 6 turns of history are maintained.")
bullet("Shift+Enter for a new line; Enter sends the message.")

heading2("4.5  Presentation — Example")
bullet("Select a large block of abstract text → '📊 Presentation'.")
body("What to verify:", bold=True)
bullet("4–8 slide cards appear, each with title and bullets.")
bullet("If AI returns invalid JSON, raw text is shown instead (graceful fallback).")

heading2("4.6  Flowchart — Example")
bullet("Right-click a methods-section PDF → '🔀 Flowchart'.")
body("What to verify:", bold=True)
bullet("Mermaid.js diagram renders in the sidebar.")
bullet("If Mermaid fails to parse, raw text is shown instead.")
note("Mermaid.js loads from CDN on first use — internet required for the first flowchart.")

heading2("4.7  Offline AI Fallback")
body("If no AI backend is available, all demo actions return a graceful message:")
code("AI backend unavailable. Please start Ollama or configure a DeepSeek API key.")

heading2("4.8  Supported File Types")
add_table(
    ["Extension", "Extraction method"],
    [
        [".pdf", "pypdf — text layer extracted page by page"],
        [".docx", "mammoth — raw text extraction"],
        [".txt / .md", "Direct file read"],
        [".html", "html2text — Markdown-like plain text"],
    ],
    col_widths=[1.5, 5.0]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5 — FILE CONVERSION
# ══════════════════════════════════════════════════════════════════
heading1("Section 5 — File Conversion")
add_table(
    ["From", "To", "Requires"],
    [
        ["PDF", "Word (.docx), plain text, HTML", "pdf2docx, pypdf (included)"],
        ["Word (.docx)", "PDF, HTML, Markdown, plain text", "mammoth (included)"],
        ["Markdown / HTML / text", "PDF, Word, HTML", "included"],
        ["Video (mp4, webm…)", "MP3, WAV, AAC, OGG, FLAC", "system ffmpeg"],
    ],
    col_widths=[2.0, 2.5, 2.0]
)

heading2("5.1  Convert PDF to Word")
bullet("Click Convert tab → select PDF → choose 'docx' → click Convert.")
body("What to verify:", bold=True)
bullet("New file '1706.03762_converted.docx' appears. Clicking it opens the Word document.")

heading2("5.2  Convert PDF to Text")
bullet("Same PDF → format 'txt' → Convert.")
body("What to verify:", bold=True)
bullet("File '1706.03762_converted.txt' appears and shows plain text in the viewer.")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6 — HISTORY & COLLECTIONS
# ══════════════════════════════════════════════════════════════════
heading1("Section 6 — History & Collections")

heading2("6.1  Viewing History")
body("What to verify:", bold=True)
bullet("Each entry shows: Filename, Source, Date, File size.")
bullet("Clicking an entry opens the file in the right viewer.")
bullet("Missing files flagged in red.")

heading2("6.2  Create a Collection")
bullet("Collections tab → 'New Collection' → enter name.")
bullet("Optionally tick 'Share with institution' (multi-user mode).")
bullet("Click Create.")
body("What to verify:", bold=True)
bullet("Collection appears in the list.")
bullet("Shared collections are accessible at GET /api/collections/shared.")

heading2("6.3  Add a Paper to a Collection")
bullet("In History, click the 🗂 button on any paper → select a collection.")
body("What to verify:", bold=True)
bullet("Opening the collection shows the paper inside it.")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 7 — CITATION EXPORT
# ══════════════════════════════════════════════════════════════════
heading1("Section 7 — Citation Export")
code("GET /api/cite/1?format=bibtex\nGET /api/cite/1?format=ris\nGET /api/cite/1?format=apa")
tip("Replace '1' with the actual history entry ID shown in the History tab.")

heading2("7.1  BibTeX")
body("What to verify:", bold=True)
bullet("Browser downloads a .bib file starting with @article{scholara:...")
bullet("Contains title, author, year, journal, url fields.")

heading2("7.2  RIS")
body("What to verify:", bold=True)
bullet("Content starts with 'TY  - JOUR' and ends with 'ER  -'.")
bullet("File imports directly into Zotero or Mendeley.")

heading2("7.3  APA")
body("What to verify:", bold=True)
bullet("Plain text in APA style: Author (Year). Title. Source. URL")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 8 — BILINGUAL INTERFACE
# ══════════════════════════════════════════════════════════════════
heading1("Section 8 — Bilingual Interface (English / French)")
body("French is the default language. Click the EN/FR button in the header, or open Settings.")
body("What to verify:", bold=True)
bullet("All tab labels, buttons, placeholders, and messages switch instantly.")
bullet("Language choice persists across page reloads.")
add_table(
    ["Element", "English", "French"],
    [
        ["Page tagline", "Open Knowledge, Everywhere", "Le Savoir Ouvert, Partout"],
        ["Search tab", "\U0001f50d Search", "\U0001f50d Recherche"],
        ["History tab", "\U0001f4cb History", "\U0001f4cb Historique"],
        ["Demo: Explanation", "Text Explanation", "Explication textuelle"],
        ["Demo: Chat", "Interactive Demo", "Démo interactive"],
        ["Demo: Flowchart", "Flowchart", "Organigramme"],
    ],
    col_widths=[2.5, 2.0, 2.0]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 9 — SETTINGS
# ══════════════════════════════════════════════════════════════════
heading1("Section 9 — Settings")
add_table(
    ["Setting", "What it does"],
    [
        ["Download directory", "Where files are saved on disk"],
        ["Max concurrent", "Parallel downloads (1–5)"],
        ["Dark mode", "Toggle dark theme"],
        ["Language", "Interface language preference"],
        ["Auto-open viewer", "Open files in viewer after download"],
    ],
    col_widths=[2.2, 4.3]
)
heading2("9.1  Check System Status")
code("http://127.0.0.1:7860/api/status")
body("What to verify:", bold=True)
bullet("JSON shows ollama_available, market_segment, active_sources, app_mode.")
bullet("institution_branding shows logo_url / primary_color when set (multi-user).")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 10 — MULTI-USER MODE
# ══════════════════════════════════════════════════════════════════
heading1("Section 10 — Multi-User Mode")
body("Multi-user mode adds JWT-based login, role management, shared institutional collections, usage analytics, and an audit log.")

heading2("10.1  Enabling Multi-User Mode")
code("APP_MODE=multi_user SECRET_KEY=your-secret python app.py")
note("Change SECRET_KEY to a long random string in production.")

heading2("10.2  User Roles")
add_table(
    ["Role", "Permissions"],
    [
        ["admin", "Full access: users, audit log, roles, impact report, shared collections"],
        ["researcher", "Search, download, convert, own history and collections"],
        ["student", "Same as researcher"],
    ],
    col_widths=[1.5, 5.0]
)

heading2("10.3  Registering and Logging In")
code("# Register\ncurl -X POST http://127.0.0.1:7860/api/auth/register \\\n  -H 'Content-Type: application/json' \\\n  -d '{\"email\":\"you@example.com\",\"password\":\"s3cr3t\"}'\n\n# Login\ncurl -X POST http://127.0.0.1:7860/api/auth/login \\\n  -H 'Content-Type: application/json' \\\n  -d '{\"email\":\"you@example.com\",\"password\":\"s3cr3t\"}'")

heading2("10.4  Admin Impact Report")
body("GET /api/admin/impact returns full analytics:")
add_table(
    ["Field", "Description"],
    [
        ["total_downloads", "Total number of files downloaded"],
        ["total_users", "Total registered users"],
        ["sources_used", "Download count per source, sorted descending"],
        ["downloads_by_day", "Download counts for the last 30 days"],
        ["top_queries", "Top 10 anonymised search query stems with counts"],
        ["top_sources", "Top 5 most-used sources"],
        ["active_users_week", "Distinct users active in the last 7 days"],
    ],
    col_widths=[2.2, 4.3]
)

heading2("10.5  Institutional Branding")
body("Set a logo_url and primary_color for your institution. The frontend applies the primary_color as a CSS variable, changing the accent colour throughout the interface.")
code("POST /api/admin/institutions\n{\"name\": \"Université de Dakar\",\n \"country\": \"SN\",\n \"logo_url\": \"https://example.edu/logo.png\",\n \"primary_color\": \"#0057a8\"}")

heading2("10.6  Shared Collections")
body("Share any collection with all users in your institution. The 'Share with institution' checkbox in the New Collection modal does this automatically.")
code("POST /api/collections/{id}/share\nGET  /api/collections/shared")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 11 — ADMIN ANALYTICS TAB
# ══════════════════════════════════════════════════════════════════
heading1("Section 11 — Admin Analytics Tab")
body("Admin users see a 📊 Admin tab in the sidebar with a live analytics panel and a bar chart of downloads by day.")
body("What to verify:", bold=True)
bullet("Total downloads and total users are shown.")
bullet("'Active this week' shows distinct users active in the last 7 days.")
bullet("Top 5 search queries with counts (anonymised query stems).")
bullet("Top sources appear with download counts.")
bullet("Bar chart shows downloads per day for the last 30 days.")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 12 — PWA & DOCKER
# ══════════════════════════════════════════════════════════════════
heading1("Section 12 — Install as an App (PWA)")

heading2("12.1  Install on Desktop")
bullet("Open Scholara in Chrome or Edge.")
bullet("Click the ⊕ icon in the address bar → 'Install Scholara'.")

heading2("12.2  Install on Android")
bullet("Phone and computer on the same Wi-Fi.")
bullet("Open Chrome on Android → http://192.168.1.x:7860")
bullet("Three-dot menu → 'Add to Home screen'.")
note("Edit the last line of app.py to use host='0.0.0.0' to serve on your local network.")

heading1("Section 13 — Running with Docker")
code("# Start with Ollama sidecar\ndocker-compose up --build\n\n# Multi-user with Docker\nAPP_MODE=multi_user SECRET_KEY=your-secret docker-compose up")
body("What to verify:", bold=True)
bullet("Scholara loads at http://localhost:7860")
bullet("/api/status shows 'ollama_available': true.")
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 14 — QUICK VERIFICATION CHECKLIST
# ══════════════════════════════════════════════════════════════════
heading1("Section 14 — Quick Verification Checklist")
add_table(
    ["#", "Feature", "Test", "Pass criterion"],
    [
        ["1",  "App loads",            "Open http://127.0.0.1:7860",               "Scholara interface appears"],
        ["2",  "Status API",           "GET /api/status",                           "JSON with tool flags"],
        ["3",  "English search",       "Search: 'climate change models'",           "arXiv / OpenAlex results"],
        ["4",  "French search",        "Search: 'changement climatique Afrique'",   "HAL results appear"],
        ["5",  "Direct download",      "URL: arxiv.org/pdf/1706.03762.pdf",         "File saved; history entry created"],
        ["6",  "Blocked URL",          "URL: google.com/file.pdf",                  "Error; no file created"],
        ["7",  "PDF → Word",       "Convert PDF to .docx",                      "New _converted.docx in list"],
        ["8",  "History view",         "Click History tab",                         "Downloads listed with date/size"],
        ["9",  "Create collection",    "Create 'My Test Collection'",               "Appears in Collections tab"],
        ["10", "Add to collection",    "Add a paper to the collection",             "Paper shows inside collection"],
        ["11", "BibTeX export",        "GET /api/cite/1?format=bibtex",             "@article{...} output returned"],
        ["12", "Language toggle",      "Click EN/FR button",                        "All labels switch instantly"],
        ["13", "Demo — explain",   "Right-click PDF → 💡 Text Explanation", "Sidebar opens with explanation"],
        ["14", "Demo — summary",   "Right-click PDF → 📄 Summary",          "Structured summary appears"],
        ["15", "Demo — chat",      "Right-click PDF → 🤖 Interactive Demo",  "Chat with follow-ups works"],
        ["16", "Demo — slides",    "Select text → 📊 Presentation",          "Slide cards rendered"],
        ["17", "Demo — flowchart", "Select text → 🔀 Flowchart",             "Mermaid diagram rendered"],
        ["18", "Demo fallback",        "Disable Ollama; try any action",            "Graceful canned message shown"],
        ["19", "Demo on selection",    "Select >20 chars → menu appears",      "Context menu shown"],
        ["20", "Search dedup",         "Search with duplicate sources",             "X-Search-Deduped header present"],
        ["21", "Admin analytics tab",  "Click 📊 Admin tab",                "Bar chart + stats render"],
        ["22", "Shared collection",    "Create collection with share ticked",       "Visible at /api/collections/shared"],
        ["23", "Institution branding", "POST institution with primary_color",       "Accent colour changes on reload"],
        ["24", "Multi-user login",     "POST /api/auth/login (multi_user mode)",    "access_token + refresh_token returned"],
        ["25", "Admin impact",         "GET /api/admin/impact (admin token)",       "downloads_by_day, top_queries, active_users_week present"],
        ["26", "PWA install",          "Open in Chrome; look for ⊕ icon",      "Install prompt appears"],
        ["27", "Queue persistence",    "Start download; restart app; check queue",  "Job re-appears as queued"],
    ],
    col_widths=[0.3, 1.5, 2.5, 2.2]
)
page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 15 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════
heading1("Section 15 — Troubleshooting")

issues = [
    ("The app does not open in the browser",
     "Make sure python app.py is running. Open http://127.0.0.1:7860 manually."),
    ("Search returns no results",
     "Check internet connection. All search sources require internet access."),
    ("AI search says 'fallback' in routing info",
     "Ollama is not running. Run 'ollama serve' and 'ollama pull mistral'."),
    ("Demo sidebar shows 'AI backend unavailable'",
     "Same as above — start Ollama or set DEEPSEEK_API_KEY. The feature degrades gracefully."),
    ("Demo sidebar does not open for my file",
     "Only .pdf, .docx, .txt, .md, and .html files are supported. Check the file extension."),
    ("Flowchart does not render",
     "Mermaid.js loads from CDN on first use — internet required. If AI returned invalid syntax, raw text is shown."),
    ("PDF conversion fails",
     "Very scanned PDFs (image-only) cannot be converted to text or Word."),
    ("Video conversion fails",
     "ffmpeg must be installed: sudo apt install ffmpeg (Linux) or brew install ffmpeg (macOS)."),
    ("Language does not switch",
     "Clear localStorage: DevTools → Application → Local Storage → delete 'lang' key."),
    ("Login returns 401 with correct credentials",
     "Set APP_MODE=multi_user and restart. In single_user mode no login endpoint is active."),
    ("Access token expired — 401 on API calls",
     "Use POST /api/auth/refresh with your refresh_token (valid 7 days) to get a new access_token."),
    ("Admin routes return 403",
     "Your account has role 'researcher'. Use PATCH /api/admin/users/{id}/role (admin token) to promote it."),
    ("Shared collections not visible",
     "Users must have the same institution_id. Check both accounts with GET /api/admin/users."),
    ("Admin analytics tab not showing",
     "The tab is hidden until an admin role is detected. In single-user mode it always shows after page load."),
    ("CORE search returns no results",
     "Set CORE_API_KEY=your-key. Without the key, CORE silently returns an empty list."),
]

for title, solution in issues:
    body(title + ":", bold=True)
    body(solution, indent=True)

# ── Back cover ────────────────────────────────────────────────────
page_break()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Scholara")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Open Knowledge, Everywhere  ·  Free & Open Source")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("arXiv · DOAJ · OpenAlex · Project Gutenberg · Internet Archive · HAL · Persée · OpenEdition · Érudit")
r.font.size = Pt(9); r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Global North: Semantic Scholar · PubMed · CrossRef · CORE · BASE")
r.font.size = Pt(9); r.font.color.rgb = GREY

doc.save("Scholara_User_Manual.docx")
print("Written: Scholara_User_Manual.docx")
